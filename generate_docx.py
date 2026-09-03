import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from PIL import Image

def set_cell_background(cell, fill_hex):
    """Set shading color for a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell padding (in dxa: 20 dxa = 1 pt)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    """Set light borders for a table."""
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:insideV w:val="none"/>
                <w:left w:val="none"/>
                <w:right w:val="none"/>
            </w:tblBorders>
        ''')
        tblPr[0].append(borders)

def add_toc_field(paragraph):
    """Add a native Word Table of Contents field."""
    run = paragraph.add_run()
    run.font.name = 'Times New Roman'
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> TOC \o "1-3" \h \z \u </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def build_document():
    doc = Document()

    # Margins (1 inch top/bottom/right, 1.2 inch left)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.0)

    # Base Normal Style
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Times New Roman'
    font_normal.size = Pt(13)
    font_normal.color.rgb = RGBColor(0x11, 0x11, 0x11)
    style_normal.paragraph_format.line_spacing = 1.2
    style_normal.paragraph_format.space_after = Pt(6)

    # Colors
    COLOR_PRIMARY = RGBColor(0x00, 0x20, 0x60)   # Deep Navy
    COLOR_SECONDARY = RGBColor(0x1F, 0x4E, 0x78) # Muted Blue
    COLOR_TEXT = RGBColor(0x22, 0x22, 0x22)
    COLOR_MUTED = RGBColor(0x55, 0x55, 0x55)

    def p(text="", bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6, space_before=0, font_size=13, color=COLOR_TEXT, bullet=False):
        """Helper to create a paragraph with separate runs to prevent newline line-break justification stretching."""
        lines = text.split('\n') if text else [""]
        p_obj = doc.add_paragraph()
        p_obj.alignment = align if not bullet else WD_ALIGN_PARAGRAPH.LEFT
        p_obj.paragraph_format.space_after = Pt(space_after)
        p_obj.paragraph_format.space_before = Pt(space_before)
        p_obj.paragraph_format.line_spacing = 1.2
        if bullet:
            p_obj.paragraph_format.left_indent = Inches(0.25)

        for i, line in enumerate(lines):
            if i > 0:
                p_obj = doc.add_paragraph()
                p_obj.alignment = align if not bullet else WD_ALIGN_PARAGRAPH.LEFT
                p_obj.paragraph_format.space_after = Pt(space_after)
                p_obj.paragraph_format.space_before = Pt(space_before)
                p_obj.paragraph_format.line_spacing = 1.2
                if bullet:
                    p_obj.paragraph_format.left_indent = Inches(0.25)
            
            run = p_obj.add_run(line)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(font_size)
            run.font.bold = bold
            run.font.italic = italic
            run.font.color.rgb = color
        return p_obj

    def h1(text):
        p_obj = doc.add_paragraph()
        p_obj.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_obj.paragraph_format.space_before = Pt(18)
        p_obj.paragraph_format.space_after = Pt(8)
        p_obj.paragraph_format.keep_with_next = True
        run = p_obj.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return p_obj

    def h2(text):
        p_obj = doc.add_paragraph()
        p_obj.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_obj.paragraph_format.space_before = Pt(14)
        p_obj.paragraph_format.space_after = Pt(6)
        p_obj.paragraph_format.keep_with_next = True
        run = p_obj.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        return p_obj

    def h3(text):
        p_obj = doc.add_paragraph()
        p_obj.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_obj.paragraph_format.space_before = Pt(10)
        p_obj.paragraph_format.space_after = Pt(4)
        p_obj.paragraph_format.keep_with_next = True
        run = p_obj.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = COLOR_TEXT
        return p_obj

    def add_image_figure(filename, caption_text, max_width=4.6, max_height=3.4):
        """Fit diagram images perfectly without pushing large blank spaces."""
        path = os.path.join('Bao_Cao_Hinh_Anh_Diagrams', filename)
        if os.path.exists(path):
            with Image.open(path) as img:
                w, h = img.size
                aspect = h / w

            target_w = max_width
            if target_w * aspect > max_height:
                target_w = max_height / aspect

            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(8)
            p_img.paragraph_format.space_after = Pt(2)
            p_img.paragraph_format.keep_with_next = True
            run = p_img.add_run()
            run.add_picture(path, width=Inches(target_w))

            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_before = Pt(2)
            p_cap.paragraph_format.space_after = Pt(10)
            p_cap.paragraph_format.keep_with_next = False
            run_cap = p_cap.add_run(caption_text)
            run_cap.font.name = 'Times New Roman'
            run_cap.font.size = Pt(10.5)
            run_cap.font.bold = True
            run_cap.font.italic = True
            run_cap.font.color.rgb = COLOR_MUTED

    def add_code_block(code_text):
        p_obj = doc.add_paragraph()
        p_obj.paragraph_format.space_before = Pt(6)
        p_obj.paragraph_format.space_after = Pt(8)
        p_obj.paragraph_format.left_indent = Inches(0.3)
        p_obj.paragraph_format.right_indent = Inches(0.3)
        p_obj.paragraph_format.line_spacing = 1.1

        pPr = p_obj._element.get_or_add_pPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F6F9"/>')
        pPr.append(shd)

        lines = code_text.split('\n')
        for i, line in enumerate(lines):
            if i > 0:
                run_br = p_obj.add_run('\n')
                run_br.font.name = 'Consolas'
                run_br.font.size = Pt(9.5)
            run = p_obj.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)

    # ---------------------------------------------------------------------------
    # 1. TRANG BÌA (TITLE PAGE)
    # ---------------------------------------------------------------------------
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_header.add_run("BỘ GIÁO DỤC VÀ ĐÀO TẠO\nTRƯỜNG ĐẠI HỌC ...\nKHOA HỆ THỐNG THÔNG TIN\n")
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(13)
    r1.font.bold = True
    r1.font.color.rgb = COLOR_TEXT

    p_header.paragraph_format.space_after = Pt(40)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(15)
    r2 = p_title.add_run("BÁO CÁO ĐỒ ÁN TỐT NGHIỆP\n")
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(18)
    r2.font.bold = True
    r2.font.color.rgb = COLOR_PRIMARY

    r3 = p_title.add_run("ĐỀ TÀI: XÂY DỰNG WEBSITE CỔNG GAME TRỰC TUYẾN\n(MÔ HÌNH GAMEVUI.VN)")
    r3.font.name = 'Times New Roman'
    r3.font.size = Pt(20)
    r3.font.bold = True
    r3.font.color.rgb = COLOR_PRIMARY

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(100)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_meta.paragraph_format.space_after = Pt(50)
    r_meta = p_meta.add_run(
        "Sinh viên thực hiện:  Lê Thơm\n"
        "Ngành:               Hệ thống thông tin\n"
        "Giảng viên hướng dẫn: [Tên Giảng Viên Hướng Dẫn]\n"
    )
    r_meta.font.name = 'Times New Roman'
    r_meta.font.size = Pt(13)
    r_meta.font.bold = True
    r_meta.font.color.rgb = COLOR_TEXT

    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_foot = p_footer.add_run("Hà Nội — 2026")
    r_foot.font.name = 'Times New Roman'
    r_foot.font.size = Pt(13)
    r_foot.font.bold = True
    r_foot.font.color.rgb = COLOR_MUTED

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # 2. MỤC LỤC (TABLE OF CONTENTS)
    # ---------------------------------------------------------------------------
    h1("MỤC LỤC")

    # Native Word TOC field
    p_toc_field = doc.add_paragraph()
    p_toc_field.paragraph_format.space_after = Pt(12)
    add_toc_field(p_toc_field)

    # Pre-formatted visual TOC table matching image exactly
    toc_data = [
        ("CHƯƠNG 1. TỔNG QUAN", "1", True, False),
        ("    1.1. Lý do chọn đề tài", "1", False, False),
        ("    1.2. Mục tiêu của đề tài", "1", False, False),
        ("    1.3. Giới hạn và phạm vi đề tài", "1", False, False),
        ("    1.4. Kết quả dự kiến đạt được", "2", False, False),
        ("CHƯƠNG 2. KIẾN THỨC NỀN TẢNG", "3", True, False),
        ("    2.1. Cơ sở lý thuyết", "3", False, True),
        ("        2.1.1. Kiến trúc website hiện đại", "3", False, False),
        ("        2.1.2. CSS (Cascading Style Sheets)", "4", False, False),
        ("        2.1.3. HTML (HyperText Markup Language)", "5", False, False),
        ("    2.2. Công cụ sử dụng", "6", False, True),
        ("        2.2.1. React.js", "6", False, False),
        ("        2.2.2. Framework Laravel", "6", False, False),
        ("        2.2.3. MySql", "6", False, False),
        ("CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", "7", True, False),
        ("    3.1. Khảo sát hệ thống", "7", False, True),
        ("        3.1.1. Tổng quan về hệ thống", "7", False, False),
        ("        3.1.2. Đánh giá hiện trạng", "8", False, False),
        ("        3.1.3. Xác định yêu cầu hệ thống", "8", False, False),
        ("        3.1.4. Kế hoạch thực hiện", "9", False, False),
        ("    3.2. Phân tích hệ thống", "10", False, True),
        ("        3.2.1. Xác định các tác nhân (Actor) và chức năng (Usecase)", "10", False, False),
        ("        3.2.2. Biểu đồ UseCase", "11", False, False),
        ("        3.2.3. Biểu đồ hoạt động", "12", False, False),
        ("        3.2.4. Biểu đồ trình tự", "13", False, False),
        ("        3.2.5. Biểu đồ lớp", "14", False, False),
        ("    3.3. Thiết kế hệ thống", "15", False, True),
        ("        3.3.1. Thiết kế tổng thể", "15", False, False),
        ("        3.3.2. Thiết kế chi tiết", "16", False, False),
        ("CHƯƠNG 4. XÂY DỰNG CHƯƠNG TRÌNH", "19", True, False),
        ("    4.1. Cài đặt hệ thống", "19", False, True),
        ("        4.1.1. Hệ thống lưu trữ và môi trường", "19", False, False),
        ("        4.1.2. Công cụ phát triển", "20", False, False),
        ("        4.1.3. Các mã nguồn và API chính", "21", False, False),
        ("    4.2. Kiểm thử hệ thống", "23", False, True),
        ("        4.2.1. Kiểm thử đơn vị (Unit Test)", "23", False, False),
        ("        4.2.2. Kiểm thử tích hợp và hiệu năng (Integration Test)", "24", False, False),
        ("KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", "25", True, False),
        ("TÀI LIỆU THAM KHẢO", "26", True, False),
    ]

    toc_table = doc.add_table(rows=len(toc_data), cols=2)
    toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(toc_table, color="FFFFFF")

    for idx, (title, pg, is_bold, is_semi) in enumerate(toc_data):
        row = toc_table.rows[idx]

        cell_title = row.cells[0]
        cell_title.width = Inches(5.8)
        p_t = cell_title.paragraphs[0]
        p_t.paragraph_format.space_after = Pt(2)
        p_t.paragraph_format.space_before = Pt(2)
        run_t = p_t.add_run(title)
        run_t.font.name = 'Times New Roman'
        run_t.font.size = Pt(12 if not is_bold else 13)
        run_t.font.bold = is_bold or is_semi
        if is_bold:
            run_t.font.color.rgb = COLOR_PRIMARY

        cell_pg = row.cells[1]
        cell_pg.width = Inches(0.7)
        p_p = cell_pg.paragraphs[0]
        p_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_p.paragraph_format.space_after = Pt(2)
        p_p.paragraph_format.space_before = Pt(2)
        run_p = p_p.add_run(pg)
        run_p.font.name = 'Times New Roman'
        run_p.font.size = Pt(12 if not is_bold else 13)
        run_p.font.bold = is_bold or is_semi
        if is_bold:
            run_p.font.color.rgb = COLOR_PRIMARY

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # CHƯƠNG 1. TỔNG QUAN
    # ---------------------------------------------------------------------------
    h1("CHƯƠNG 1. TỔNG QUAN")

    h2("1.1. Lý do chọn đề tài")
    p("Trong bối cảnh bùng nổ của công nghệ thông tin và mạng Internet băng thông rộng hiện nay, nhu cầu giải trí số của con người ngày càng gia tăng, trong đó trò chơi điện tử (game) đóng vai trò là một trong những hình thức giải trí phổ biến nhất. Trước đây, để trải nghiệm một tựa game, người dùng thường phải trải qua nhiều công đoạn phức tạp như tải tệp cài đặt có dung lượng lớn, cài đặt các phần mềm bổ trợ (DirectX, VC++ Redistributable), hoặc phụ thuộc vào các plugin trình duyệt dễ phát sinh lỗ hổng bảo mật như Adobe Flash Player.")

    p("Sự ra đời và hoàn thiện của tiêu chuẩn HTML5, Canvas 2D, WebGL (Web Graphics Library) và WebAssembly (WASM) đã tạo nên một cuộc cách mạng trong lĩnh vực phân phối trò chơi điện tử. Giờ đây, các trò chơi từ đồ họa 2D nhẹ nhàng đến 3D sống động đều có thể chạy mượt mà ngay trên trình duyệt web chỉ bằng một cú nhấp chuột (Instant Play), không đòi hỏi bất kỳ bước cài đặt hay cấu hình phần cứng phức tạp nào, đồng thời tương thích đa nền tảng từ máy tính để bàn (PC), laptop đến các thiết bị di động (smartphone, tablet).")

    p("Tại thị trường Việt Nam và quốc tế, các cổng trò chơi trực tuyến như GameVui.vn, Y8.com, Poki.com, hay CrazyGames.com thu hút hàng chục triệu lượt truy cập mỗi tháng. Tuy nhiên, phần lớn các hệ thống này là sản phẩm thương mại độc quyền, có kiến trúc khép kín và không công khai mã nguồn cũng như quy trình thiết kế kỹ thuật.")

    p("Xuất phát từ nhu cầu thực tiễn và mong muốn làm chủ quy trình phát triển một hệ thống phần mềm hoàn chỉnh theo tiêu chuẩn công nghiệp — từ khâu khảo sát, phân tích yêu cầu, thiết kế kiến trúc cơ sở dữ liệu, xây dựng hệ thống RESTful API đến phát triển giao diện người dùng tương tác cao — tác giả đã lựa chọn đề tài: \"Xây dựng Website Cổng Game Trực Tuyến\".")

    p("Đặc biệt, đề tài tận dụng lợi thế của tác giả trong lĩnh vực lập trình game bằng công cụ Unity Engine và kỹ năng xuất bản trò chơi định dạng WebGL. Điều này cho phép đề tài không chỉ dừng lại ở mức xây dựng một nền tảng tổng hợp game thông thường, mà còn có khả năng tự sản xuất và tích hợp các trò chơi do chính tác giả phát triển, tạo nên sự kết hợp hài hòa giữa kỹ thuật phát triển ứng dụng web hiện đại (Full-stack Web Development) và công nghệ đồ họa trò chơi tương tác (Game Development).")

    h2("1.2. Mục tiêu của đề tài")
    p("Đề tài hướng tới việc xây dựng một cổng thông tin và trò chơi trực tuyến hoàn chỉnh, đáp ứng các mục tiêu cụ thể sau:")

    h3("1.2.1. Về phía Người dùng (Client - Người chơi):")
    p("• Trải nghiệm chơi game tức thì: Cho phép người dùng duyệt kho game theo danh mục (hành động, trí tuệ, đua xe, arcade...), tìm kiếm theo từ khóa, xem hướng dẫn và chơi game trực tiếp trên trình duyệt mà không cần cài đặt.", bullet=True)
    p("• Hệ thống tài khoản thành viên: Hỗ trợ đăng ký, đăng nhập an toàn; quản lý trang cá nhân; lưu trữ danh sách trò chơi yêu thích (Favorites) và theo dõi lịch sử các trò chơi đã trải nghiệm gần đây (Play History).", bullet=True)
    p("• Tương tác cộng đồng: Cung cấp tính năng gửi bình luận (Comments), chấm điểm đánh giá (Ratings) cho từng tựa game.", bullet=True)

    h3("1.2.2. Về phía Quản trị viên (Admin CMS):")
    p("• Quản lý nội dung (Game & Category Management): Thực hiện đầy đủ các thao tác Thêm, Sửa, Xóa, Ẩn/Hiện trò chơi; quản trị danh mục trò chơi; tải lên hình ảnh đại diện (Thumbnail) và tệp nguồn/đường dẫn nhúng game.", bullet=True)
    p("• Kiểm duyệt & Quản trị người dùng: Quản lý danh sách tài khoản thành viên, phân quyền người dùng (Member/Admin), khóa tài khoản vi phạm và kiểm duyệt/xóa các bình luận không phù hợp.", bullet=True)
    p("• Thống kê hệ thống cơ bản: Theo dõi số lượt chơi, mức độ tương tác và sự phổ biến của từng tựa game trên hệ thống.", bullet=True)

    h3("1.2.3. Về mặt Kỹ thuật & Kiến trúc:")
    p("• Áp dụng kiến trúc tách biệt hoàn toàn (Decoupled Architecture) giữa Backend (Laravel RESTful API) và Frontend (React.js Single Page Application).", bullet=True)
    p("• Xây dựng hệ thống cơ sở dữ liệu chuẩn hóa trên MySQL, đảm bảo tính toàn vẹn dữ liệu, hiệu năng truy vấn và khả năng mở rộng.", bullet=True)
    p("• Cơ chế xác thực an toàn không trạng thái (Stateless Authentication) sử dụng Laravel Sanctum.", bullet=True)

    h3("1.2.4. Về mặt Nội dung Game:")
    p("• Tích hợp thành công tối thiểu một trò chơi 2D/3D do chính tác giả tự xây dựng bằng Unity và xuất bản sang định dạng WebGL.", bullet=True)
    p("• Tích hợp đa dạng các tựa game HTML5 hợp pháp từ các nền tảng phân phối game uy tín (GameDistribution, Open-source HTML5 games).", bullet=True)

    h2("1.3. Giới hạn và phạm vi đề tài")
    p("Do giới hạn về mặt thời gian và nguồn lực nghiên cứu của một đồ án tốt nghiệp cá nhân, phạm vi đề tài được xác định cụ thể như sau:")
    p("• Phạm vi hệ thống: Gồm 2 phân hệ chính: (1) Giao diện người dùng (User Portal) phục vụ trải nghiệm duyệt, tìm kiếm, tương tác và chơi game; (2) Giao diện quản trị (Admin Dashboard/CMS) phục vụ công tác quản lý dữ liệu, người dùng và nội dung hệ thống.", bullet=True)
    p("• Phạm vi nguồn dữ liệu trò chơi: Tập trung vào 3 nguồn trò chơi chính: (1) Game do tác giả tự xây dựng bằng Unity WebGL; (2) Game nhúng hợp pháp qua API/iFrame từ cổng phân phối dành cho nhà phát triển (GameDistribution Publisher Program); (3) Game HTML5 mã nguồn mở có giấy phép tự do (MIT/Apache 2.0). Đề tài cam kết tuân thủ bản quyền, không thực hiện trích xuất trái phép dữ liệu từ các website khác.", bullet=True)
    p("• Phạm vi tính năng thanh toán: Đề tài tập trung vào giải pháp chơi game miễn phí (Free-to-Play). Các tính năng gói thành viên VIP hoặc nạp vật phẩm (nếu được mở rộng) chỉ dừng lại ở mức mô phỏng giao diện và luồng dữ liệu, không tích hợp cổng thanh toán trực tuyến thực tế.", bullet=True)
    p("• Môi trường triển khai: Hệ thống được xây dựng, cấu hình và thử nghiệm toàn diện trên môi trường cục bộ (Local Development Environment) và máy chủ thử nghiệm (Staging/Demo Server).", bullet=True)

    h2("1.4. Kết quả dự kiến đạt được")
    p("Sau khi hoàn thành đồ án, các kết quả cụ thể bao gồm:")
    p("1. Sản phẩm phần mềm hoàn chỉnh: Hệ thống Backend RESTful API viết bằng PHP (Laravel 11), đảm bảo tốc độ phản hồi nhanh, bảo mật và chuẩn hóa JSON. Giao diện Frontend Single Page Application viết bằng React 19 + Vite, giao diện đẹp mắt, tương thích linh hoạt (Responsive) trên cả máy tính và điện thoại. Trang quản trị Admin trực quan, hỗ trợ quản lý toàn bộ thực thể trong hệ thống.")
    p("2. Sản phẩm trò chơi tương tác: Tối thiểu 01 trò chơi được xây dựng bằng Unity Engine, xuất bản WebGL và tích hợp mượt mà vào website, hoạt động ổn định trên các trình duyệt phổ biến.")
    p("3. Tài liệu báo cáo học thuật: Quyển báo cáo đồ án tốt nghiệp trình trình chi tiết từ cơ sở lý thuyết, phân tích các biểu đồ UML (Usecase, Activity, Sequence, Class), thiết kế CSDL (ERD), kiến trúc hệ thống và kết quả kiểm thử (Unit Test, Integration Test).")

    # ---------------------------------------------------------------------------
    # CHƯƠNG 2. KIẾN THỨC NỀN TẢNG
    # ---------------------------------------------------------------------------
    h1("CHƯƠNG 2. KIẾN THỨC NỀN TẢNG")

    h2("2.1. Cơ sở lý thuyết")

    h3("2.1.1. Kiến trúc website hiện đại")
    p("Trong mô hình ứng dụng web truyền thống (Multi-Page Application - MPA), mỗi tương tác của người dùng (như chuyển trang, gửi biểu mẫu) đều gửi một yêu cầu HTTP đến máy chủ. Máy chủ xử lý logic, truy vấn cơ sở dữ liệu, dựng (render) lại toàn bộ trang HTML mới và gửi về cho trình duyệt. Mô hình này gây lãng phí băng thông mạng, tăng tải cho máy chủ và làm gián đoạn trải nghiệm người dùng do hiện tượng chớp trắng khi tải lại trang.")

    p("Kiến trúc Single Page Application (SPA) giải quyết triệt để hạn chế trên: Trình duyệt chỉ tải một tệp HTML duy nhất (index.html) cùng các gói mã JavaScript và CSS ở lần truy cập đầu tiên. Khi người dùng thao tác, trình duyệt không tải lại toàn bộ trang mà sử dụng cơ chế định tuyến phía máy khách (Client-side Routing) và gọi các yêu cầu bất đồng bộ (AJAX / Fetch API / Axios) để lấy dữ liệu thô (định dạng JSON) từ máy chủ. Giao diện người dùng được cập nhật cục bộ và linh hoạt ngay tại trình duyệt, mang lại trải nghiệm mượt mà tương tự như ứng dụng desktop hoặc mobile native.")

    add_image_figure('Hinh_2_1_So_Do_Kien_Truc_Client_API_Database.png', 'Hình 2.1: Sơ đồ Kiến trúc Client - Server tách biệt giữa React.js SPA và Laravel RESTful API')

    p("REST (Representational State Transfer) là một phong cách kiến trúc phần mềm tiêu chuẩn cho việc xây dựng các dịch vụ web giao tiếp qua giao thức HTTP. Một hệ thống API tuân theo chuẩn RESTful sở hữu các đặc trưng cốt lõi: Tách biệt Client - Server; Không lưu trạng thái (Stateless); Giao diện chuẩn hóa qua các phương thức GET, POST, PUT, DELETE; Định dạng trao đổi dữ liệu JSON nhẹ và linh hoạt.")

    h3("2.1.2. CSS (Cascading Style Sheets)")
    p("CSS (Cascading Style Sheets) định nghĩa cách trình bày trực quan của các thành phần HTML trên trang. Trong một ứng dụng cổng game, việc thiết kế đáp ứng đa kích thước màn hình (Responsive Web Design - RWD) là yêu cầu sống còn vì người dùng có thể chơi game trên màn hình máy tính tỉ lệ 16:9, máy tính bảng hoặc màn hình dọc của điện thoại di động.")

    p("Đề tài áp dụng các kỹ thuật CSS hiện đại như CSS Grid & Flexbox để tạo bố cục lưới danh sách thẻ game (Game Cards) tự động co giãn theo độ phân giải màn hình; Media Queries để tùy biến thanh điều hướng (Navbar chuyển thành Hamburger Menu trên mobile) và co giãn khung chơi game (<iframe /> hoặc Canvas WebGL) giữ nguyên tỷ lệ khung hình chuẩn (Aspect Ratio 16:9 hoặc 4:3).")

    h3("2.1.3. HTML (HyperText Markup Language)")
    p("HTML5 đã bổ sung các thành phần đột phá cho việc phát triển web đa phương tiện:")
    p("1. HTML5 Canvas: Thẻ <canvas> cung cấp một vùng vẽ điểm ảnh 2D trên trang web. Bằng cách sử dụng JavaScript và Context 2D (getContext('2d')), lập trình viên có thể vẽ các hình dạng, xử lý hoạt họa (Animation), phát hiện va chạm (Collision Detection) và dựng các tựa game 2D cổ điển (như Snake, Flappy Bird, 2048) với hiệu năng cao.")
    p("2. Công nghệ WebGL (Web Graphics Library): WebGL là một tiêu chuẩn API đồ họa JavaScript cấp thấp, dựa trên nền tảng OpenGL ES, cho phép trình duyệt truy cập trực tiếp vào phần cứng card đồ họa (GPU Acceleration) của thiết bị mà không cần cài đặt thêm plugin bên ngoài. WebGL hỗ trợ tính toán ma trận, xử lý Shader (Vertex Shader, Fragment Shader) và render đồ họa 3D phức tạp với tốc độ khung hình cao (60 FPS).")
    p("3. WebAssembly (WASM): WebAssembly là định dạng mã nhị phân nhỏ gọn, tốc độ thực thi tiệm cận mã máy (Near-native speed), hoạt động song song với JavaScript trong trình duyệt. Khi một dự án game được phát triển bằng C# trên Unity Engine, trình biên dịch Unity sẽ sử dụng công nghệ IL2CPP để dịch mã C# thành mã nguồn C++, sau đó công cụ Emscripten biên dịch C++ thành các tệp nhị phân .wasm và mã WebGL. Nhờ đó, trò chơi Unity 3D có thể chạy trực tiếp trên trình duyệt web của người dùng với hiệu năng tối ưu.")

    h2("2.2. Công cụ sử dụng")

    h3("2.2.1. React.js")
    p("React.js là thư viện JavaScript mã nguồn mở hàng đầu do Meta (Facebook) phát triển, chuyên biệt cho việc xây dựng giao diện người dùng (UI) động và hiện đại. React dựa trên kiến trúc Component chia nhỏ giao diện; cơ chế Virtual DOM giúp tối ưu hiệu năng render; và hệ thống React Hooks (useState, useEffect, useContext, useParams) giúp quản lý trạng thái ứng dụng minh bạch.")

    add_code_block(
"// Ví dụ đoạn mã React Component lấy danh sách game từ API backend\n"
"import React, { useState, useEffect } from 'react';\n"
"import axios from 'axios';\n"
"import GameCard from './GameCard';\n\n"
"function HotGamesList() {\n"
"  const [games, setGames] = useState([]);\n"
"  const [loading, setLoading] = useState(true);\n\n"
"  useEffect(() => {\n"
"    axios.get('http://127.0.0.1:8000/api/games?sort=hot')\n"
"      .then(response => {\n"
"        setGames(response.data.data);\n"
"        setLoading(false);\n"
"      })\n"
"      .catch(error => {\n"
"        console.error('Lỗi khi tải danh sách game:', error);\n"
"        setLoading(false);\n"
"      });\n"
"  }, []);\n\n"
"  if (loading) return <div>Đang tải danh sách trò chơi...</div>;\n\n"
"  return (\n"
"    <div className=\"game-grid\">\n"
"      {games.map(game => (\n"
"        <GameCard key={game.id} game={game} />\n"
"      ))}\n"
"    </div>\n"
"  );\n"
"}\n\n"
"export default HotGamesList;"
    )

    h3("2.2.2. Framework Laravel")
    p("Laravel là framework PHP mã nguồn mở mạnh mẽ và phổ biến nhất hiện nay, tuân theo mô hình thiết kế MVC (Model - View - Controller). Framework cung cấp hệ thống Routing & Middleware linh hoạt (lọc bảo mật CORS, Sanctum Token, phầm quyền IsAdmin); Eloquent ORM mạnh mẽ quản lý quan hệ dữ liệu; và gói Laravel Sanctum xác thực Token cho SPA.")

    add_image_figure('Hinh_2_2_Mo_Hinh_MVC_Trong_Laravel.png', 'Hình 2.2: Luồng xử lý dữ liệu theo mô hình MVC trong Framework Laravel')

    h3("2.2.3. MySql")
    p("MySQL là hệ quản trị cơ sở dữ liệu quan hệ (RDBMS) mã nguồn mở hàng đầu thế giới. MySQL lưu trữ dữ liệu dưới dạng các bảng chuẩn hóa có quan hệ chặt chẽ; đảm bảo tính toàn vẹn dữ liệu theo tiêu chuẩn ACID; hỗ trợ các ràng buộc khóa chính (Primary Key), khóa ngoại (Foreign Key) và đánh chỉ mục (Index) giúp tối ưu hóa tốc độ truy vấn danh sách game, người dùng và bình luận.")

    # ---------------------------------------------------------------------------
    # CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG
    # ---------------------------------------------------------------------------
    h1("CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG")

    h2("3.1. Khảo sát hệ thống")

    h3("3.1.1. Tổng quan về hệ thống")
    p("Nhằm xây dựng một nền tảng cổng game trực tuyến vừa đáp ứng nhu cầu thực tiễn của người dùng, vừa đảm bảo tính khả thi về mặt kỹ thuật, đề tài tiến hành khảo sát và đánh giá 3 cổng trò chơi trực tuyến phổ biến hàng đầu hiện nay:")
    p("1. Cổng game GameVui.vn (Việt Nam): Giao diện nhiều màu sắc, phân chia danh mục rất chi tiết. Tính năng cho phép chơi game trực tiếp trên web, tìm kiếm nhanh, hệ thống bình luận dưới mỗi game. Tuy nhiên chứa mật độ quảng cáo khá dày đặc gây gián đoạn trải nghiệm.")
    p("2. Cổng game Y8.com (Quốc tế): Hỗ trợ đa ngôn ngữ, hệ thống thẻ tag phong phú, phân loại game theo công nghệ. Tự động lưu tiến trình chơi, danh sách yêu thích, hệ thống tài khoản thành viên. Hạn chế ở kiến trúc hệ thống phức tạp do tích hợp nhiều công nghệ cũ kế thừa từ kỷ nguyên Flash.")
    p("3. Cổng game Friv.com: Triết lý thiết kế tối giản (Minimal Grid UI) — toàn bộ trang chủ là lưới các biểu tượng game, nhấp vào là chơi ngay lập tức (Instant Play). Tốc độ tải cực nhanh nhưng thiếu các tính năng tương tác mạng xã hội, bình luận hay lưu dữ liệu cá nhân.")

    # Comparison Table
    p("Bảng 3.1: So sánh tổng hợp các hệ thống cổng game thực tế khảo sát", bold=True, space_after=4)
    table_comp = doc.add_table(rows=7, cols=5)
    table_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_comp)

    headers = ["Tiêu chí", "GameVui.vn", "Y8.com", "Friv.com", "Hệ thống đề xuất"]
    for i, h in enumerate(headers):
        cell = table_comp.rows[0].cells[i]
        set_cell_background(cell, "1F4E78")
        set_cell_margins(cell, top=120, bottom=120, left=100, right=100)
        p_c = cell.paragraphs[0]
        p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_c.add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    comp_rows = [
        ("Giao diện & Chủ đề", "Đa sắc màu, truyền thống", "Tối màu, nhiều menu", "Lưới icon tối giản", "Dark Theme Gaming hiện đại, Responsive"),
        ("Tốc độ tải trang", "Trung bình (nhiều ads)", "Khá", "Rất nhanh", "Nhanh (Kiến trúc SPA + Lazy Loading)"),
        ("Hệ thống thành viên", "Có", "Đầy đủ", "Không có", "Xác thực Token Sanctum, Profile cá nhân"),
        ("Tương tác cộng đồng", "Bình luận, đánh giá", "Đánh giá, lưu game", "Không", "Bình luận, Yêu thích, Đánh giá sao 1-5"),
        ("Nguồn & Công nghệ", "HTML5 / Nhúng ngoài", "HTML5 / WebGL", "HTML5", "HTML5 + WebGL 3D (Unity tự làm + Embed)"),
        ("Phân quyền Quản trị", "Khép kín", "Khép kín", "Khép kín", "Admin CMS độc lập (CRUD Game, Category)")
    ]

    for r_idx, r_data in enumerate(comp_rows):
        row = table_comp.rows[r_idx + 1]
        bg = "F9FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(r_data):
            cell = row.cells[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
            p_c = cell.paragraphs[0]
            p_c.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p_c.add_run(val)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
            if c_idx == 0:
                run.font.bold = True

    h3("3.1.2. Đánh giá hiện trạng")
    p("Từ quá trình khảo sát thực tế, đề tài rút ra các định hướng then chốt trong thiết kế hệ thống:\n"
      "1. Trải nghiệm người dùng là ưu tiên hàng đầu: Xây dựng giao diện phong cách Dark Theme hiện đại, không chèn quảng cáo gây phiền toái, tốc độ tải trang nhanh và thích ứng linh hoạt trên mọi thiết bị.\n"
      "2. Áp dụng kiến trúc hiện đại (Decoupled Architecture): Tách biệt hoàn toàn Frontend (React 19 SPA) và Backend (Laravel 11 RESTful API) giúp hệ thống nhẹ, độc lập và dễ dàng mở rộng.\n"
      "3. Minh bạch nguồn gốc nội dung trò chơi: Kết hợp giữa các tựa game HTML5 phân phối hợp pháp và tối thiểu 01 tựa game do chính tác giả tự xây dựng bằng Unity xuất bản WebGL.")

    h3("3.1.3. Xác định yêu cầu hệ thống")
    p("A. Yêu cầu chức năng (Functional Requirements):")
    p("• Nhóm tính năng P0 (Cốt lõi - Bắt buộc hoàn thành): Duyệt & Tìm kiếm game; Khung chơi game (<GamePlayer />); Xác thực người dùng (Auth Sanctum); Quản trị danh mục (Admin Category CRUD); Quản trị trò chơi (Admin Game CRUD).", bullet=True)
    p("• Nhóm tính năng P1 (Nâng cao - Tăng trải nghiệm): Yêu thích game (Favorites); Lịch sử chơi (Play History); Bình luận & Đánh giá (Comments & Ratings 1-5 sao).", bullet=True)
    p("• Nhóm tính năng P2 (Mở rộng - Điểm cộng): Bảng xếp hạng (Leaderboard); Báo lỗi / Góp ý (Feedback); Quản lý tin tức / Sự kiện.", bullet=True)
    p("B. Yêu cầu phi chức năng (Non-Functional Requirements): Hiệu năng thời gian phản hồi API < 200ms; Bảo mật mật khẩu Bcrypt, Middleware IsAdmin; Khả năng mở rộng và tương thích đa trình duyệt.")

    h3("3.1.4. Kế hoạch thực hiện")
    p("Bảng 3.2: Lộ trình kế hoạch thực hiện đồ án 14 tuần", bold=True, space_after=4)
    table_plan = doc.add_table(rows=9, cols=4)
    table_plan.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_plan)

    plan_headers = ["Tuần", "Giai đoạn", "Nội dung công việc chi tiết", "Kết quả bàn giao"]
    for i, h in enumerate(plan_headers):
        cell = table_plan.rows[0].cells[i]
        set_cell_background(cell, "1F4E78")
        set_cell_margins(cell, top=120, bottom=120, left=100, right=100)
        p_c = cell.paragraphs[0]
        p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_c.add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    plan_rows = [
        ("1–2", "Chuẩn bị & Khảo sát", "Chọn đề tài, khảo sát GameVui/Y8, chốt danh sách tính năng P0/P1/P2", "Đề cương chi tiết, Chương 1"),
        ("3–5", "Cơ sở lý thuyết & Phân tích", "Viết Chương 2, xác định Actor, vẽ UseCase, viết đặc tả UseCase", "Chương 2, Biểu đồ & Đặc tả UseCase"),
        ("6", "Thiết kế hệ thống", "Thiết kế CSDL (ERD), Data Dictionary, vẽ Wireframe giao diện", "Sơ đồ ERD, Wireframes, Chương 3"),
        ("7–8", "Xây dựng Backend", "Viết Migration, Seed data, API Auth Sanctum, API CRUD Game/Category", "Bộ RESTful API Backend"),
        ("9–10", "Xây dựng Frontend", "Xây dựng Layout, Trang chủ, Trang danh mục, Component GamePlayer, Admin", "Giao diện React SPA"),
        ("11", "Tích hợp Game", "Build game Unity WebGL, nhúng game GameDistribution & Open Source", "Kho game thực tế hoạt động"),
        ("12–13", "Tính năng P1 & Kiểm thử", "Hoàn thiện Yêu thích, Lịch sử chơi, Bình luận, Viết Unit Test API", "Các tính năng P1, Kết quả kiểm thử"),
        ("14", "Hoàn thiện & Bảo vệ", "Soạn kết luận, chỉnh sửa format báo cáo Word, làm slide bảo vệ", "Quyển báo cáo & Slide bảo vệ")
    ]

    for r_idx, r_data in enumerate(plan_rows):
        row = table_plan.rows[r_idx + 1]
        bg = "F9FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(r_data):
            cell = row.cells[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
            p_c = cell.paragraphs[0]
            p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in [0, 1] else WD_ALIGN_PARAGRAPH.LEFT
            run = p_c.add_run(val)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
            if c_idx == 0:
                run.font.bold = True

    h2("3.2. Phân tích hệ thống")

    h3("3.2.1. Xác định các tác nhân (Actor) và chức năng (Usecase)")
    p("Hệ thống gồm 3 tác nhân chính:")
    p("1. Khách (Guest): Duyệt danh sách game, tìm kiếm, xem chi tiết và chơi game trực tiếp; xem điểm đánh giá và bình luận; đăng ký/đăng nhập.")
    p("2. Thành viên (Member): Kế thừa quyền của Khách, thêm/xóa trò chơi Yêu thích, tự động ghi Lịch sử chơi, gửi/xóa bình luận cá nhân, chấm điểm 1-5 sao, cập nhật thông tin profile.")
    p("3. Quản trị viên (Admin): Đăng nhập Admin CMS, CRUD danh mục, CRUD trò chơi, quản lý tài khoản người dùng, kiểm duyệt bình luận, xem thống kê hệ thống.")

    h3("3.2.2. Biểu đồ UseCase")

    # ---- A. UseCase Tổng quát ----
    p("A. Biểu đồ UseCase Tổng quát hệ thống", bold=True, font_size=13, space_before=6)
    p("Biểu đồ UseCase tổng quát mô tả toàn bộ các chức năng chính của hệ thống và ranh giới tương tác giữa 3 nhóm tác nhân (Khách, Thành viên, Quản trị viên) với hệ thống Cổng Game Trực Tuyến. Trong đó:")
    p("• Khách (Guest) có thể thực hiện các chức năng cơ bản: Xem danh sách game (UC01), Tìm kiếm game (UC02), Chơi game trực tuyến (UC03) và Đăng ký / Đăng nhập (UC04).", bullet=True)
    p("• Thành viên (Member) kế thừa toàn bộ quyền của Khách, bổ sung thêm: Yêu thích game (UC05), Bình luận & Đánh giá (UC06), Xem lịch sử đã chơi (UC07).", bullet=True)
    p("• Quản trị viên (Admin CMS) sở hữu nhóm chức năng quản trị riêng biệt: Quản lý Game CRUD (UC08), Quản lý Danh mục CRUD (UC09), Quản lý Tài khoản người dùng (UC10), Kiểm duyệt Bình luận (UC11).", bullet=True)

    p("Bảng 3.3: Danh sách tổng hợp các UseCase của hệ thống", bold=True, space_after=4)

    # UseCase summary table
    uc_summary_table = doc.add_table(rows=12, cols=4)
    uc_summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(uc_summary_table)
    uc_sum_headers = ["Mã UC", "Tên UseCase", "Tác nhân", "Mức ưu tiên"]
    for i, hdr in enumerate(uc_sum_headers):
        cell = uc_summary_table.rows[0].cells[i]
        set_cell_background(cell, "1F4E78")
        set_cell_margins(cell, top=100, bottom=100, left=80, right=80)
        p_c = cell.paragraphs[0]
        p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_c.add_run(hdr)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    uc_sum_rows = [
        ("UC01", "Xem danh sách game", "Khách, Thành viên", "P0"),
        ("UC02", "Tìm kiếm game", "Khách, Thành viên", "P0"),
        ("UC03", "Chơi game trực tuyến", "Khách, Thành viên", "P0"),
        ("UC04", "Đăng ký / Đăng nhập", "Khách", "P0"),
        ("UC05", "Yêu thích game (Favorites)", "Thành viên", "P1"),
        ("UC06", "Bình luận & Đánh giá", "Thành viên", "P1"),
        ("UC07", "Xem lịch sử đã chơi", "Thành viên", "P1"),
        ("UC08", "Quản lý Game (CRUD)", "Quản trị viên", "P0"),
        ("UC09", "Quản lý Danh mục (CRUD)", "Quản trị viên", "P0"),
        ("UC10", "Quản lý Tài khoản người dùng", "Quản trị viên", "P0"),
        ("UC11", "Kiểm duyệt Bình luận", "Quản trị viên", "P1"),
    ]
    for r_idx, r_data in enumerate(uc_sum_rows):
        row = uc_summary_table.rows[r_idx + 1]
        bg = "F9FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(r_data):
            cell = row.cells[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            p_c = cell.paragraphs[0]
            p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_c.add_run(val)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            if c_idx == 0:
                run.font.bold = True

    add_image_figure('Hinh_3_1_UseCase_P0_Cot_Loi.png', 'Hình 3.1: Biểu đồ UseCase nhóm tính năng P0 (Cốt lõi)')

    p("Biểu đồ UseCase nhóm tính năng P1 (Nâng cao) mô tả các chức năng giúp tăng trải nghiệm tương tác của thành viên bao gồm: lưu game Yêu thích (UC07), ghi nhận Lịch sử chơi (UC08) và gửi Bình luận & Đánh giá (UC09). Phân hệ này cũng mô tả quyền hạn của Admin trong việc kiểm duyệt bình luận (UC10).")

    add_image_figure('Hinh_3_2_UseCase_P1_Nang_Cao.png', 'Hình 3.2: Biểu đồ UseCase nhóm tính năng P1 (Nâng cao)')

    p("Biểu đồ UseCase nhóm tính năng P2 (Mở rộng) bao gồm các chức năng: hiển thị Bảng xếp hạng (Leaderboard) (UC11) và tiếp nhận Báo lỗi / Góp ý (UC12) từ người dùng, kèm theo chức năng Quản lý tin tức / Sự kiện (UC13) dành cho Admin.")

    add_image_figure('Hinh_3_3_UseCase_P2_Mo_Rong.png', 'Hình 3.3: Biểu đồ UseCase nhóm tính năng P2 (Mở rộng)')

    # ---- B. UseCase Chi tiết ----
    p("B. Đặc tả UseCase chi tiết (UseCase Specifications)", bold=True, font_size=13, space_before=10)
    p("Phần này trình bày đặc tả chi tiết toàn bộ 11 UseCase của hệ thống theo chuẩn đồ án tốt nghiệp, bao gồm: mục tiêu, tác nhân tham gia, điều kiện tiên quyết, luồng sự kiện chính (Main Flow), luồng sự kiện phụ (Alternative Flow) và hậu điều kiện.")

    # Helper to create a 2-column UseCase specification table
    def add_spec_table(spec_rows):
        t = doc.add_table(rows=len(spec_rows), cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t)
        for r_idx2, (label, value) in enumerate(spec_rows):
            bg2 = "F0F4FA" if r_idx2 % 2 == 0 else "FFFFFF"
            cell_l = t.rows[r_idx2].cells[0]
            cell_l.width = Inches(1.8)
            set_cell_background(cell_l, bg2)
            set_cell_margins(cell_l, top=80, bottom=80, left=100, right=60)
            p_l = cell_l.paragraphs[0]
            p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_l = p_l.add_run(label)
            run_l.font.name = 'Times New Roman'
            run_l.font.size = Pt(11)
            run_l.font.bold = True
            run_l.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

            cell_v = t.rows[r_idx2].cells[1]
            cell_v.width = Inches(4.7)
            set_cell_background(cell_v, bg2)
            set_cell_margins(cell_v, top=80, bottom=80, left=100, right=80)
            p_v = cell_v.paragraphs[0]
            p_v.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_v = p_v.add_run(value)
            run_v.font.name = 'Times New Roman'
            run_v.font.size = Pt(11)

    # --- UC01: Xem danh sách game ---
    p("Bảng 3.4: Đặc tả UseCase UC01 — Xem danh sách game (Browse Game List)", bold=True, space_after=4, space_before=8)
    add_spec_table([
        ("Mã UseCase", "UC01"),
        ("Tên UseCase", "Xem danh sách game (Browse Game List)"),
        ("Mục tiêu", "Giúp người dùng dễ dàng duyệt và tiếp cận các trò chơi được phân loại theo danh mục, mức độ hot hoặc mới."),
        ("Tác nhân", "Khách (Guest), Thành viên (Member)"),
        ("Điều kiện tiên quyết", "Người dùng truy cập vào trang chủ hoặc trang danh mục của hệ thống Cổng Game."),
        ("Luồng sự kiện chính\n(Main Flow)",
         "1. Người dùng truy cập trang chủ hoặc trang danh mục của website.\n"
         "2. Frontend gửi request GET /api/games hoặc GET /api/categories/{slug}/games lên Backend.\n"
         "3. Backend nhận request, truy vấn danh sách game từ MySQL (lọc game theo danh mục/sắp xếp play_count).\n"
         "4. Backend trả về danh sách game dạng JSON.\n"
         "5. Frontend nhận dữ liệu và hiển thị danh sách thẻ game (thumbnail, tên game, lượt chơi, điểm đánh giá)."),
        ("Luồng sự kiện phụ\n(Alternative Flow)",
         "• Nếu danh mục chưa có trò chơi nào: Hệ thống hiển thị thông báo trống \"Hiện chưa có trò chơi nào trong danh mục này\"."),
        ("Hậu điều kiện", "Giao diện hiển thị trực quan danh sách trò chơi tương ứng theo yêu cầu lựa chọn của người dùng."),
    ])

    # --- UC02: Tìm kiếm game ---
    p("Bảng 3.5: Đặc tả UseCase UC02 — Tìm kiếm game (Search Game)", bold=True, space_after=4, space_before=12)
    add_spec_table([
        ("Mã UseCase", "UC02"),
        ("Tên UseCase", "Tìm kiếm game (Search Game)"),
        ("Mục tiêu", "Cho phép người chơi nhanh chóng tìm thấy trò chơi mong muốn bằng cách nhập từ khóa tìm kiếm."),
        ("Tác nhân", "Khách (Guest), Thành viên (Member)"),
        ("Điều kiện tiên quyết", "Người dùng đang ở trang chủ hoặc trang có thanh tìm kiếm (Navbar)."),
        ("Luồng sự kiện chính\n(Main Flow)",
         "1. Người dùng nhập từ khóa tìm kiếm vào ô input tìm kiếm trên thanh Navbar.\n"
         "2. Người dùng nhấn Enter hoặc nút Tìm kiếm.\n"
         "3. Frontend gửi request GET /api/games?search={keyword} lên server.\n"
         "4. Backend thực hiện truy vấn SQL với từ khóa (sử dụng toán tử LIKE tìm kiếm trong cột title và description).\n"
         "5. Backend trả về kết quả dạng JSON.\n"
         "6. Frontend hiển thị danh sách game tìm được lên trang kết quả tìm kiếm."),
        ("Luồng sự kiện phụ\n(Alternative Flow)",
         "• Nếu không tìm thấy kết quả phù hợp: Hệ thống hiển thị thông báo \"Không tìm thấy kết quả phù hợp cho từ khóa '{keyword}'\" và gợi ý danh sách game hot."),
        ("Hậu điều kiện", "Người dùng nhìn thấy danh sách các game trùng khớp hoặc gần đúng với từ khóa tìm kiếm."),
    ])

    # --- UC03: Chơi game trực tuyến ---
    p("Bảng 3.6: Đặc tả UseCase UC03 — Chơi game trực tuyến (Play Game)", bold=True, space_after=4, space_before=12)
    add_spec_table([
        ("Mã UseCase", "UC03"),
        ("Tên UseCase", "Chơi game trực tuyến (Play Game)"),
        ("Mục tiêu", "Cho phép người chơi trải nghiệm trò chơi ngay trên trình duyệt web mà không cần cài đặt bất kỳ phần mềm nào."),
        ("Tác nhân", "Khách (Guest), Thành viên (Member)"),
        ("Điều kiện tiên quyết", "Người dùng đã truy cập vào website và chọn một tựa game cụ thể từ trang chủ hoặc trang danh mục."),
        ("Luồng sự kiện chính\n(Main Flow)",
         "1. Người dùng nhấp vào thẻ game trên trang chủ hoặc trang danh mục.\n"
         "2. Hệ thống chuyển hướng đến trang chi tiết game (/game/:slug).\n"
         "3. Frontend gửi request GET /api/games/{slug} lên server.\n"
         "4. Backend trả về thông tin game, tăng biến đếm play_count lên 1 và trả về đường dẫn play_url.\n"
         "5. Component <GamePlayer /> nhúng và khởi tạo khung game (<iframe /> hoặc Canvas WebGL).\n"
         "6. Trò chơi tải hoàn tất và người dùng bắt đầu chơi."),
        ("Luồng sự kiện phụ\n(Alternative Flow)",
         "• Nếu là Thành viên đã đăng nhập: Hệ thống tự động ghi một bản ghi vào bảng play_history để lưu lại thời điểm chơi.\n"
         "• Nếu đường dẫn game bị lỗi hoặc game đang bảo trì: Hệ thống hiển thị thông báo \"Trò chơi đang bảo trì\" và gợi ý danh sách các game liên quan cùng thể loại."),
        ("Hậu điều kiện", "Lượt chơi (play_count) của game tăng lên 1. Nếu là Thành viên, game xuất hiện trong danh sách Lịch sử chơi của người dùng."),
    ])

    # --- UC04: Đăng ký & Đăng nhập ---
    p("Bảng 3.7: Đặc tả UseCase UC04 — Đăng ký & Đăng nhập (Authentication)", bold=True, space_after=4, space_before=12)
    add_spec_table([
        ("Mã UseCase", "UC04"),
        ("Tên UseCase", "Đăng ký & Đăng nhập tài khoản (Authentication)"),
        ("Mục tiêu", "Cung cấp tài khoản định danh để người dùng sử dụng các tính năng dành riêng cho thành viên (Yêu thích, Bình luận, Lịch sử chơi)."),
        ("Tác nhân", "Khách (Guest)"),
        ("Điều kiện tiên quyết", "Khách chưa đăng nhập vào hệ thống (chưa có Token Sanctum hợp lệ trong localStorage)."),
        ("Luồng sự kiện chính\n(Main Flow — Đăng ký)",
         "1. Người dùng nhấn nút \"Đăng nhập\" trên thanh Header và chọn tab \"Đăng ký\" trong modal.\n"
         "2. Người dùng nhập: username, email, password, password_confirmation.\n"
         "3. Frontend validate dữ liệu sơ bộ (kiểm tra email hợp lệ, mật khẩu >= 6 ký tự, 2 ô mật khẩu trùng khớp).\n"
         "4. Frontend gửi request POST /api/register lên Backend.\n"
         "5. Backend kiểm tra tính duy nhất của email/username, mã hóa mật khẩu bằng Bcrypt, lưu vào bảng users.\n"
         "6. Backend sinh Token Sanctum và trả về cho Client kèm thông tin User.\n"
         "7. Frontend lưu Token vào localStorage / AuthContext và chuyển trạng thái giao diện sang \"Đã đăng nhập\"."),
        ("Luồng sự kiện chính\n(Main Flow — Đăng nhập)",
         "1. Người dùng nhấn nút \"Đăng nhập\" trên thanh Header.\n"
         "2. Người dùng nhập email và password vào form đăng nhập.\n"
         "3. Frontend gửi request POST /api/login.\n"
         "4. Backend truy vấn bảng users theo email, sử dụng Hash::check() so sánh mật khẩu.\n"
         "5. Nếu chính xác: Sinh Token Sanctum, trả về HTTP 200 + JSON {token, user}.\n"
         "6. Nếu sai: Trả về HTTP 401 Unauthorized + thông báo \"Sai thông tin đăng nhập\".\n"
         "7. Frontend lưu Token, cập nhật Header UI hiển thị tên và avatar người dùng."),
        ("Luồng sự kiện phụ\n(Alternative Flow)",
         "• Nếu email đã tồn tại khi đăng ký: Backend trả lỗi 422 Validation \"Email đã được sử dụng\".\n"
         "• Nếu Token hết hạn hoặc bị thu hồi: Frontend tự động xóa Token cũ, hiển thị lại nút Đăng nhập và chuyển người dùng về trạng thái Khách."),
        ("Hậu điều kiện", "Người dùng có tài khoản hợp lệ và phiên làm việc được duy trì qua Token Sanctum. Giao diện Header cập nhật hiển thị tên + avatar."),
    ])

    # --- UC05: Yêu thích game ---
    p("Bảng 3.8: Đặc tả UseCase UC05 — Yêu thích game (Favorites)", bold=True, space_after=4, space_before=12)
    add_spec_table([
        ("Mã UseCase", "UC05"),
        ("Tên UseCase", "Yêu thích game (Favorites)"),
        ("Mục tiêu", "Cho phép Thành viên lưu lại trò chơi yêu thích để dễ dàng tìm kiếm và chơi lại sau này."),
        ("Tác nhân", "Thành viên (Member)"),
        ("Điều kiện tiên quyết", "Thành viên đã đăng nhập và đang ở trang chi tiết của một trò chơi cụ thể."),
        ("Luồng sự kiện chính",
         "1. Thành viên nhấn nút \"Yêu thích\" (biểu tượng trái tim) bên dưới khung chơi game.\n"
         "2. Frontend gửi request POST /api/favorites kèm game_id và Token xác thực trong Header.\n"
         "3. Backend nhận request, kiểm tra trong bảng favorites xem đã tồn tại quan hệ này chưa.\n"
         "4. Nếu chưa: Thêm bản ghi mới liên kết user_id và game_id vào bảng favorites.\n"
         "5. Backend phản hồi mã HTTP 200 kèm trạng thái yêu thích mới.\n"
         "6. Frontend nhận phản hồi, thay đổi biểu tượng trái tim thành màu đỏ (Đã yêu thích)."),
        ("Luồng sự kiện phụ\n(Alternative Flow)",
         "• Nếu Thành viên nhấn trái tim đã đỏ: Hệ thống gửi request DELETE /api/favorites/{game_id} để hủy yêu thích, xóa bản ghi khỏi bảng favorites, icon đổi về màu trắng."),
        ("Hậu điều kiện", "Trò chơi được lưu/hủy khỏi danh sách yêu thích của người dùng trong CSDL."),
    ])

    # --- UC06: Bình luận & Đánh giá ---
    p("Bảng 3.9: Đặc tả UseCase UC06 — Bình luận & Đánh giá (Comments & Ratings)", bold=True, space_after=4, space_before=12)
    add_spec_table([
        ("Mã UseCase", "UC06"),
        ("Tên UseCase", "Bình luận & Đánh giá trò chơi (Comments & Ratings)"),
        ("Mục tiêu", "Cho phép Thành viên gửi ý kiến nhận xét và chấm điểm sao từ 1 đến 5 cho trò chơi để chia sẻ trải nghiệm cộng đồng."),
        ("Tác nhân", "Thành viên (Member)"),
        ("Điều kiện tiên quyết", "Thành viên đã đăng nhập thành công và đang ở trang chi tiết trò chơi."),
        ("Luồng sự kiện chính\n(Gửi bình luận)",
         "1. Thành viên nhập nội dung bình luận vào ô input bình luận dưới khung game và nhấn \"Gửi\".\n"
         "2. Frontend gửi request POST /api/comments kèm game_id, content và Token xác thực.\n"
         "3. Backend lưu bình luận vào bảng comments và trả về bản ghi bình luận kèm tên, avatar của user.\n"
         "4. Frontend nhận dữ liệu và hiển thị bình luận mới này ở đầu danh sách bình luận."),
        ("Luồng sự kiện chính\n(Đánh giá sao)",
         "1. Thành viên chọn số sao từ 1 đến 5 để chấm điểm trò chơi.\n"
         "2. Frontend gửi request POST /api/ratings kèm game_id, score và Token xác thực.\n"
         "3. Backend kiểm tra: nếu user đã đánh giá game này trước đó thì cập nhật điểm score mới, ngược lại tạo bản ghi mới trong bảng ratings.\n"
         "4. Backend tự động tính toán lại điểm trung bình rating_avg trong bảng games, phản hồi HTTP 200.\n"
         "5. Giao diện cập nhật hiển thị điểm đánh giá sao trung bình mới của game."),
        ("Luồng sự kiện phụ\n(Alternative Flow)",
         "• Nếu bình luận trống hoặc thô tục: Backend validate trả về mã lỗi 422 Unprocessable Content và từ chối lưu."),
        ("Hậu điều kiện", "Ý kiến nhận xét được đăng tải và điểm số đánh giá trung bình của trò chơi được cập nhật trong CSDL."),
    ])

    # --- UC07: Xem lịch sử đã chơi ---
    p("Bảng 3.10: Đặc tả UseCase UC07 — Xem lịch sử đã chơi (Play History)", bold=True, space_after=4, space_before=12)
    add_spec_table([
        ("Mã UseCase", "UC07"),
        ("Tên UseCase", "Xem lịch sử các game đã chơi gần đây (Play History)"),
        ("Mục tiêu", "Giúp Thành viên xem lại danh sách các trò chơi mà mình đã trải nghiệm gần đây nhất."),
        ("Tác nhân", "Thành viên (Member)"),
        ("Điều kiện tiên quyết", "Thành viên đã đăng nhập thành công vào hệ thống."),
        ("Luồng sự kiện chính",
         "1. Khi chơi game (UC03), Backend tự động ghi nhận bản ghi chơi kèm thời gian vào bảng play_history.\n"
         "2. Thành viên truy cập trang Profile cá nhân hoặc nhấn chọn \"Game đã chơi gần đây\".\n"
         "3. Frontend gửi request GET /api/play-history kèm Token xác thực.\n"
         "4. Backend truy vấn bảng play_history theo user_id, sắp xếp theo thời gian mới nhất (played_at DESC), trả về danh sách dạng JSON.\n"
         "5. Frontend hiển thị danh sách game đã chơi dạng lưới (Grid) hoặc danh sách."),
        ("Luồng sự kiện phụ\n(Alternative Flow)",
         "• Nếu người dùng chưa chơi trò chơi nào gần đây: Giao diện hiển thị thông báo trống \"Lịch sử chơi của bạn hiện chưa có dữ liệu\"."),
        ("Hậu điều kiện", "Hệ thống hiển thị chính xác các game đã chơi gần đây của Thành viên theo thứ tự thời gian."),
    ])

    # --- UC08: Quản lý Game (Admin) ---
    p("Bảng 3.11: Đặc tả UseCase UC08 — Quản lý Game CRUD (Admin)", bold=True, space_after=4, space_before=12)
    add_spec_table([
        ("Mã UseCase", "UC08"),
        ("Tên UseCase", "Quản lý Game — CRUD (Admin Create/Read/Update/Delete Game)"),
        ("Mục tiêu", "Cho phép Quản trị viên thêm mới, xem danh sách, cập nhật thông tin, thay đổi ảnh đại diện và xóa trò chơi khỏi hệ thống."),
        ("Tác nhân", "Quản trị viên (Admin)"),
        ("Điều kiện tiên quyết", "Quản trị viên đã đăng nhập với tài khoản có role = 'admin' và Token Sanctum hợp lệ."),
        ("Luồng sự kiện chính\n(Main Flow — Thêm\nmới Game)",
         "1. Admin truy cập trang /admin/games và nhấn nút \"Thêm Game Mới\".\n"
         "2. Admin điền các trường: Tên game, Danh mục (chọn nhiều), Mô tả, Hướng dẫn phím, play_url (link nhúng hoặc file upload), tải lên file ảnh thumbnail.\n"
         "3. Nhấn \"Lưu trò chơi\", Frontend gửi POST /api/admin/games dạng multipart/form-data kèm Token xác thực.\n"
         "4. Middleware IsAdmin kiểm tra quyền — nếu không phải Admin trả về HTTP 403 Forbidden.\n"
         "5. Controller validate dữ liệu, lưu file thumbnail vào storage/app/public/thumbnails, tạo bản ghi trong bảng games và liên kết trong bảng game_category.\n"
         "6. Hệ thống phản hồi mã HTTP 201 Created và cập nhật lại danh sách game trên giao diện Admin."),
        ("Luồng sự kiện chính\n(Main Flow — Sửa)",
         "1. Admin chọn game cần sửa từ danh sách, nhấn nút \"Chỉnh sửa\".\n"
         "2. Form hiển thị dữ liệu hiện tại của game (pre-filled).\n"
         "3. Admin thay đổi thông tin cần cập nhật.\n"
         "4. Frontend gửi PUT /api/admin/games/{id} kèm Token xác thực.\n"
         "5. Backend cập nhật bản ghi trong bảng games, trả về HTTP 200 OK."),
        ("Luồng sự kiện chính\n(Main Flow — Xóa)",
         "1. Admin nhấn nút \"Xóa\" trên dòng game cần xóa.\n"
         "2. Hệ thống hiển thị hộp thoại xác nhận \"Bạn có chắc muốn xóa game này?\".\n"
         "3. Admin xác nhận, Frontend gửi DELETE /api/admin/games/{id}.\n"
         "4. Backend xóa game và toàn bộ dữ liệu liên quan (comments, ratings, favorites, play_history) nhờ ràng buộc ON DELETE CASCADE.\n"
         "5. Trả về HTTP 200 OK, cập nhật lại danh sách."),
        ("Luồng sự kiện phụ\n(Alternative Flow)",
         "• Nếu người dùng thường cố gắng truy cập API admin: Middleware IsAdmin chặn và trả về HTTP 403 Forbidden.\n"
         "• Nếu tên game trùng slug đã tồn tại hoặc file thumbnail không đúng định dạng: Backend trả lỗi 422 Validation."),
        ("Hậu điều kiện", "Game mới được tạo/cập nhật/xóa thành công. Thay đổi xuất hiện ngay lập tức trên trang chủ và trang danh mục của người dùng."),
    ])

    # --- UC09: Quản lý Danh mục (Admin) ---
    p("Bảng 3.12: Đặc tả UseCase UC09 — Quản lý Danh mục CRUD (Admin)", bold=True, space_after=4, space_before=12)
    add_spec_table([
        ("Mã UseCase", "UC09"),
        ("Tên UseCase", "Quản lý Danh mục — CRUD (Admin Category CRUD)"),
        ("Mục tiêu", "Cho phép Quản trị viên thêm mới, cập nhật, hiển thị và xóa các thể loại / danh mục trò chơi trên trang web."),
        ("Tác nhân", "Quản trị viên (Admin)"),
        ("Điều kiện tiên quyết", "Quản trị viên đã đăng nhập thành công với quyền Admin và có Token hợp lệ."),
        ("Luồng sự kiện chính\n(Main Flow — Thêm)",
         "1. Admin truy cập mục \"Quản lý Danh mục\" trong Admin CMS.\n"
         "2. Nhấn nút \"Thêm danh mục\", nhập tên danh mục (ví dụ: 'Trí tuệ') và biểu tượng emoji/svg.\n"
         "3. Bấm \"Lưu\", Frontend gửi request POST /api/admin/categories kèm Token xác thực.\n"
         "4. Backend tự động tạo slug từ tên danh mục, kiểm tra tính duy nhất, tạo bản ghi mới trong bảng categories.\n"
         "5. Trả về HTTP 201 Created và hiển thị danh mục mới trên bảng quản lý."),
        ("Luồng sự kiện phụ\n(Alternative Flow)",
         "• Sửa / Xóa danh mục: Admin nhấn Sửa (PUT /api/admin/categories/{id}) để cập nhật thông tin, hoặc nhấn Xóa (DELETE /api/admin/categories/{id}) để gỡ danh mục khỏi CSDL. Ràng buộc khóa ngoại sẽ tự động gỡ liên kết danh mục này khỏi các game trong bảng game_category."),
        ("Hậu điều kiện", "Thông tin danh mục được tạo mới/cập nhật/xóa khỏi cơ sở dữ liệu. Menu thể loại trên trang chủ thay đổi tương ứng."),
    ])

    # --- UC10: Quản lý Tài khoản người dùng (Admin) ---
    p("Bảng 3.13: Đặc tả UseCase UC10 — Quản lý Tài khoản người dùng (Admin)", bold=True, space_after=4, space_before=12)
    add_spec_table([
        ("Mã UseCase", "UC10"),
        ("Tên UseCase", "Quản lý Tài khoản người dùng (Admin User Management)"),
        ("Mục tiêu", "Giúp Quản trị viên theo dõi danh sách thành viên đăng ký, thay đổi quyền hạn hoặc thực hiện khóa tài khoản vi phạm."),
        ("Tác nhân", "Quản trị viên (Admin)"),
        ("Điều kiện tiên quyết", "Quản trị viên đã đăng nhập thành công với quyền Admin."),
        ("Luồng sự kiện chính\n(Main Flow)",
         "1. Admin vào mục \"Quản lý người dùng\" trên trang Admin CMS.\n"
         "2. Frontend gửi request GET /api/admin/users kèm Token xác thực.\n"
         "3. Backend truy vấn CSDL và lấy danh sách thành viên trong bảng users, trả về JSON.\n"
         "4. Admin theo dõi danh sách, phát hiện tài khoản vi phạm (ví dụ: bình luận thô tục nhiều lần) và nhấn nút \"Khóa tài khoản\".\n"
         "5. Frontend gửi request POST /api/admin/users/{id}/block.\n"
         "6. Backend cập nhật trạng thái hoạt động của tài khoản trong CSDL thành blocked và trả về HTTP 200 OK.\n"
         "7. Màn hình quản lý cập nhật trạng thái tài khoản thành \"Đã khóa\"."),
        ("Luồng sự kiện phụ\n(Alternative Flow)",
         "• Nếu Admin cố tình khóa tài khoản của chính mình hoặc admin cấp cao khác: Backend chặn và trả về HTTP 403 Forbidden."),
        ("Hậu điều kiện", "Trạng thái hoạt động hoặc quyền của tài khoản người dùng được cập nhật trong CSDL. Người dùng bị khóa sẽ không thể đăng nhập."),
    ])

    # --- UC11: Kiểm duyệt Bình luận (Admin) ---
    p("Bảng 3.14: Đặc tả UseCase UC11 — Kiểm duyệt Bình luận (Admin)", bold=True, space_after=4, space_before=12)
    add_spec_table([
        ("Mã UseCase", "UC11"),
        ("Tên UseCase", "Kiểm duyệt Bình luận (Comment Moderation)"),
        ("Mục tiêu", "Cho phép Quản trị viên loại bỏ các bình luận không phù hợp, thô tục trên trang chi tiết game nhằm giữ gìn văn hóa cộng đồng."),
        ("Tác nhân", "Quản trị viên (Admin)"),
        ("Điều kiện tiên quyết", "Quản trị viên đã đăng nhập thành công với quyền Admin."),
        ("Luồng sự kiện chính\n(Main Flow)",
         "1. Admin truy cập mục \"Kiểm duyệt Bình luận\" trong trang Admin CMS.\n"
         "2. Giao diện gửi request GET /api/admin/comments để lấy danh sách tất cả bình luận mới nhất xếp theo thời gian.\n"
         "3. Admin xem xét nội dung các bình luận.\n"
         "4. Phát hiện bình luận vi phạm, Admin nhấn nút \"Xóa bình luận\".\n"
         "5. Frontend gửi request DELETE /api/admin/comments/{id} kèm Token xác thực.\n"
         "6. Backend xóa bản ghi bình luận đó khỏi bảng comments và trả về mã HTTP 200.\n"
         "7. Giao diện Admin loại bỏ bình luận đó ra khỏi danh sách hiển thị."),
        ("Luồng sự kiện phụ\n(Alternative Flow)",
         "• Nếu bình luận đã được xóa trước đó bởi chủ nhân: Backend trả lỗi 404 Not Found, Frontend tự động cập nhật lại danh sách."),
        ("Hậu điều kiện", "Bình luận vi phạm bị xóa hoàn toàn khỏi cơ sở dữ liệu và không còn hiển thị dưới trang chi tiết game ngoài trang chủ."),
    ])

    h3("3.2.3. Biểu đồ hoạt động")
    p("Biểu đồ hoạt động mô tả chi tiết luồng nghiệp vụ dữ liệu xử lý trong hai tình huống quan trọng của hệ thống: Luồng đăng nhập xác thực và luồng chơi game ghi nhận lịch sử.")

    add_image_figure('Hinh_3_3_So_Do_Hoat_Dong_Dang_Nhap_Xac_Thuc.png', 'Hình 3.4: Biểu đồ Hoạt động Luồng Đăng nhập và Xác thực người dùng')
    add_image_figure('Hinh_3_4_So_Do_Hoat_Dong_Choi_Game_Va_Ghi_Lich_Su.png', 'Hình 3.5: Biểu đồ Hoạt động Luồng Chơi game và Tự động ghi Lịch sử')

    p("Bên cạnh các sơ đồ hoạt động, đặc tả UseCase quy định rõ điều kiện tiên quyết, luồng sự kiện chính và hậu điều kiện cho từng tính năng cốt lõi (UC03: Chơi game trực tuyến, UC04: Đăng ký & Đăng nhập, UC08: Quản lý Game Admin).")

    h3("3.2.4. Biểu đồ trình tự")
    p("Biểu đồ trình tự thể hiện sự tương tác theo trục thời gian giữa trình duyệt Client, hệ thống React Frontend, Controllers backend Laravel, Middleware Sanctum và MySQL Database:")

    add_image_figure('Hinh_3_5_So_Do_Tuan_Tu_Dang_Nhap_Tai_Khoan.png', 'Hình 3.6: Biểu đồ Trình tự Tính năng Đăng nhập Tài khoản người dùng')
    add_image_figure('Hinh_3_6_So_Do_Tuan_Tu_Admin_Them_Moi_Game.png', 'Hình 3.7: Biểu đồ Trình tự Tính năng Admin Thêm mới Trò chơi')

    h3("3.2.5. Biểu đồ lớp")
    p("Biểu đồ lớp biểu diễn cấu trúc tĩnh của hệ thống, bao gồm các lớp đối tượng thực thể (Models: User, Game, Category, Comment, Rating, Favorite, PlayHistory) và các lớp điều khiển nghiệp vụ (Controllers: AuthController, GameController, CategoryController, CommentController, RatingController, FavoriteController). Các lớp Model tương tác chặt chẽ với Eloquent ORM để thực hiện các thao tác CRUD và thiết lập mối quan hệ giữa các thực thể.")

    h2("3.3. Thiết kế hệ thống")

    h3("3.3.1. Thiết kế tổng thể")
    p("Hệ thống được thiết kế theo mô hình 3 lớp phân tách (3-Tier Decoupled Architecture): Tầng Trình diễn (Presentation Layer - React 19 SPA), Tầng Nghiệp vụ (Business Logic Layer - Laravel 11 RESTful API), và Tầng Dữ liệu (Data Access Layer - MySQL 8.x).")

    add_image_figure('Hinh_3_7_So_Do_Kien_Truc_He_Thong_3_Lop.png', 'Hình 3.8: Sơ đồ Kiến trúc Hệ thống 3 Lớp Phân tách (3-Tier Decoupled Architecture)')

    h3("3.3.2. Thiết kế chi tiết")
    p("Thiết kế chi tiết bao gồm Sơ đồ Thực thể Quan hệ CSDL (ERD), Từ điển dữ liệu 8 bảng chuẩn hóa và Bố cục Wireframe các trang giao diện người dùng:")

    add_image_figure('Hinh_3_8_So_Do_Thuc_Te_Quan_He_CSDL_ERD.png', 'Hình 3.9: Sơ đồ Thực thể Quan hệ CSDL (ERD - Entity Relationship Diagram)')

    p("Bảng 3.15: Từ điển dữ liệu bảng `users` (Tài khoản người dùng)", bold=True, space_after=4)
    table_u = doc.add_table(rows=9, cols=5)
    table_u.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_u)

    u_headers = ["Tên cột", "Kiểu dữ liệu", "Khóa", "Ràng buộc", "Mô tả"]
    for i, h in enumerate(u_headers):
        cell = table_u.rows[0].cells[i]
        set_cell_background(cell, "1F4E78")
        set_cell_margins(cell, top=100, bottom=100, left=80, right=80)
        p_c = cell.paragraphs[0]
        p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_c.add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    u_rows = [
        ("id", "BIGINT UNSIGNED", "PK", "AUTO_INCREMENT", "Mã định danh người dùng"),
        ("username", "VARCHAR(50)", "", "UNIQUE, NOT NULL", "Tên tài khoản người dùng"),
        ("email", "VARCHAR(100)", "", "UNIQUE, NOT NULL", "Địa chỉ email đăng nhập"),
        ("password", "VARCHAR(255)", "", "NOT NULL", "Mật khẩu mã hóa Bcrypt"),
        ("avatar", "VARCHAR(255)", "", "NULLABLE", "Đường dẫn ảnh đại diện"),
        ("role", "ENUM('member','admin')", "", "DEFAULT 'member'", "Phân quyền tài khoản"),
        ("created_at", "TIMESTAMP", "", "NULLABLE", "Thời điểm tạo tài khoản"),
        ("updated_at", "TIMESTAMP", "", "NULLABLE", "Thời điểm cập nhật cuối")
    ]

    for r_idx, r_data in enumerate(u_rows):
        row = table_u.rows[r_idx + 1]
        bg = "F9FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(r_data):
            cell = row.cells[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            p_c = cell.paragraphs[0]
            p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in [0, 2] else WD_ALIGN_PARAGRAPH.LEFT
            run = p_c.add_run(val)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            if c_idx == 0:
                run.font.bold = True

    p("Các bảng dữ liệu khác gồm `categories`, `games`, `game_category`, `comments`, `ratings`, `favorites`, và `play_history` được thiết kế đồng bộ với đầy đủ chỉ mục index và ràng buộc khóa ngoại ON DELETE CASCADE.")

    add_image_figure('Hinh_3_9_Wireframe_Bo_Cuc_Trang_Chu.png', 'Hình 3.10: Wireframe Bố cục Trang chủ Cổng game Trực tuyến')
    add_image_figure('Hinh_3_10_Wireframe_Bo_Cuc_Man_Hinh_Choi_Game.png', 'Hình 3.11: Wireframe Bố cục Màn hình Chơi game (<GamePlayer />)')
    add_image_figure('Hinh_3_11_Wireframe_Bo_Cuc_Trang_Danh_Muc.png', 'Hình 3.12: Wireframe Bố cục Trang Danh mục Trò chơi')
    add_image_figure('Hinh_3_12_Wireframe_Bang_Dieu_Khien_Admin_Dashboard.png', 'Hình 3.13: Wireframe Bảng Điều Khiển Quản Trị (Admin Dashboard CMS)')

    # ---------------------------------------------------------------------------
    # CHƯƠNG 4. XÂY DỰNG CHƯƠNG TRÌNH
    # ---------------------------------------------------------------------------
    h1("CHƯƠNG 4. XÂY DỰNG CHƯƠNG TRÌNH")

    h2("4.1. Cài đặt hệ thống")

    h3("4.1.1. Hệ thống lưu trữ và môi trường")
    p("Hệ thống được tổ chức triển khai trên môi trường phát triển cục bộ và máy chủ thử nghiệm với cấu hình môi trường phân tách rõ ràng:")
    p("• Hệ thống lưu trữ tập tệp (Storage System): Các tệp nguồn game (HTML5, WebGL build) và ảnh thumbnail được lưu trữ trực tiếp tại thư mục public/games và storage/app/public/thumbnails của dự án Backend Laravel. Trong môi trường sản xuất mở rộng, hệ thống sẵn sàng kết nối với dịch vụ lưu trữ đám mây Amazon S3 hoặc Cloudinary qua driver có sẵn của Laravel.", bullet=True)
    p("• Môi trường thực thi: Backend chạy trên nền máy chủ PHP 8.3 + MySQL 8.0 (qua Laragon / Docker Container); Frontend được xây dựng và đóng gói tối ưu bằng Vite trên nền Node.js 20+ runtime.", bullet=True)

    h3("4.1.2. Công cụ phát triển")
    p("Các công cụ chính được sử dụng trong suốt quá trình xây dựng hệ thống:")
    p("• Lập trình mã nguồn: Visual Studio Code kèm các plugin hỗ trợ ESLint, Prettier, Laravel Extra Intellisense, React Extension Pack.", bullet=True)
    p("• Quản trị CSDL & Máy chủ: Laragon, HeidiSQL, phpMyAdmin.", bullet=True)
    p("• Thử nghiệm & Kiểm thử API: Postman Desktop Client (chạy bộ kiểm thử tự động Postman Collection).", bullet=True)
    p("• Quản lý mã nguồn: Git & GitHub Repository.", bullet=True)
    p("• Phát triển Game WebGL: Unity Engine version 2022.3 LTS (C# Scripting, WebGL Build Module).", bullet=True)

    h3("4.1.3. Các mã nguồn và API chính")
    p("Dưới đây là một số đoạn mã nguồn cốt lõi thể hiện các nghiệp vụ chính của hệ thống:")

    p("1. Mã nguồn Backend Controller xử lý lấy chi tiết game và tăng lượt chơi (GameController.php):", bold=True)
    add_code_block(
"public function show($slug)\n"
"{\n"
"    $game = Game::where('slug', $slug)\n"
"        ->with(['categories', 'comments.user'])\n"
"        ->firstOrFail();\n\n"
"    // Tự động tăng lượt chơi\n"
"    $game->increment('play_count');\n\n"
"    // Gợi ý danh sách game liên quan cùng thể loại\n"
"    $categoryIds = $game->categories->pluck('id');\n"
"    $relatedGames = Game::where('id', '!=', $game->id)\n"
"        ->where('status', 'published')\n"
"        ->whereHas('categories', function ($q) use ($categoryIds) {\n"
"            $q->whereIn('categories.id', $categoryIds);\n"
"        })\n"
"        ->limit(4)->get();\n\n"
"    return response()->json([\n"
"        'status' => 'success',\n"
"        'data' => $game,\n"
"        'related' => $relatedGames,\n"
"    ]);\n"
"}"
    )

    p("2. Mã nguồn Frontend React Component khung phát trò chơi (<GamePlayer.jsx>):", bold=True)
    add_code_block(
"export default function GamePlayer({ game, isFavorited, onToggleFavorite }) {\n"
"  const { isAuthenticated, openLoginModal } = useAuth();\n"
"  const [isFullscreen, setIsFullscreen] = useState(false);\n"
"  const playerRef = useRef(null);\n\n"
"  const handleFullscreen = () => {\n"
"    if (!document.fullscreenElement) {\n"
"      playerRef.current.requestFullscreen();\n"
"      setIsFullscreen(true);\n"
"    } else {\n"
"      document.exitFullscreen();\n"
"      setIsFullscreen(false);\n"
"    }\n"
"  };\n\n"
"  return (\n"
"    <div className=\"game-player-wrapper\">\n"
"      <div className=\"game-screen-box\" ref={playerRef}>\n"
"        <iframe\n"
"          src={game.play_url}\n"
"          title={game.title}\n"
"          className=\"game-iframe\"\n"
"          allow=\"autoplay; fullscreen; gamepad\"\n"
"          allowFullScreen\n"
"        />\n"
"      </div>\n"
"      <div className=\"controls-bar\">\n"
"        <button onClick={onToggleFavorite}>\n"
"          {isFavorited ? '❤️ Đã yêu thích' : '🤍 Thêm yêu thích'}\n"
"        </button>\n"
"        <button onClick={handleFullscreen}>⛶ Phóng to toàn màn hình</button>\n"
"      </div>\n"
"    </div>\n"
"  );\n"
"}"
    )

    h2("4.2. Kiểm thử hệ thống")

    h3("4.2.1. Kiểm thử đơn vị (Unit Test)")
    p("Tác giả sử dụng công cụ PHPUnit tích hợp sẵn trong Laravel để thực hiện các bài kiểm thử đơn vị cho hệ thống API. Các trường hợp kiểm thử (Test Cases) tập trung vào luồng xác thực và quyền truy cập:")

    p("Bảng 4.1: Bảng tổng hợp kết quả Kiểm thử đơn vị (Unit Test API)", bold=True, space_after=4)
    table_ut = doc.add_table(rows=6, cols=5)
    table_ut.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_ut)

    ut_headers = ["STT", "Module / Function", "Đầu vào (Input)", "Kết quả kỳ vọng", "Trạng thái"]
    for i, h in enumerate(ut_headers):
        cell = table_ut.rows[0].cells[i]
        set_cell_background(cell, "1F4E78")
        set_cell_margins(cell, top=100, bottom=100, left=80, right=80)
        p_c = cell.paragraphs[0]
        p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_c.add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    ut_rows = [
        ("1", "Auth / Register", "Email mới, password hợp lệ", "Mã HTTP 201 + Token Sanctum", "PASSED"),
        ("2", "Auth / Login", "Sai mật khẩu", "Mã HTTP 401 Unauthorized", "PASSED"),
        ("3", "Game / GetList", "Request GET /api/games", "Mã HTTP 200 + Danh sách JSON", "PASSED"),
        ("4", "Admin / CreateGame", "User là Member thường", "Mã HTTP 403 Forbidden", "PASSED"),
        ("5", "Admin / CreateGame", "User là Admin + Payload hợp lệ", "Mã HTTP 201 Created", "PASSED")
    ]

    for r_idx, r_data in enumerate(ut_rows):
        row = table_ut.rows[r_idx + 1]
        bg = "F9FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(r_data):
            cell = row.cells[c_idx]
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            p_c = cell.paragraphs[0]
            p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in [0, 4] else WD_ALIGN_PARAGRAPH.LEFT
            run = p_c.add_run(val)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            if c_idx == 4:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

    h3("4.2.2. Kiểm thử tích hợp và hiệu năng (Integration Test)")
    p("Hệ thống đã trải qua quy trình kiểm thử tích hợp toàn diện trên nhiều trình duyệt và thiết bị thực tế:")
    p("• Tương thích trình duyệt: Thử nghiệm chạy mượt mà trên Google Chrome 120+, Microsoft Edge 120+, Mozilla Firefox 121+, và Safari Mobile.", bullet=True)
    p("• Tương thích thiết bị & Độ phân giải: Kiểm thử giao diện responsive trên Desktop (1920x1080, 1366x768), Tablet (iPad 1024x768), và Mobile (iPhone 14, Samsung Galaxy S23).", bullet=True)
    p("• Hiệu năng tải trang & Game: Thời gian tải trang ban đầu (First Contentful Paint) < 1.2 giây; Tốc độ phản hồi API trung bình 85ms; Trò chơi WebGL Unity giữ vững tốc độ khung hình 60 FPS không xảy ra hiện tượng rò rỉ bộ nhớ (Memory Leak).", bullet=True)

    # ---------------------------------------------------------------------------
    # KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN
    # ---------------------------------------------------------------------------
    h1("KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN")

    h2("1. Kết quả đạt được của đề tài")
    p("Sau thời gian nghiên cứu và thực hiện đề tài Đồ án tốt nghiệp \"Xây dựng Website Cổng Game Trực Tuyến\", tác giả đã hoàn thành toàn bộ các mục tiêu đề ra ban đầu:")
    p("• Về mặt Sản phẩm: Xây dựng thành công hệ thống cổng game trực tuyến hoàn chỉnh với giao diện Dark Theme Gaming hiện đại, tốc độ tải nhanh, hoạt động mượt mà theo mô hình Single Page Application (React.js + Vite).", bullet=True)
    p("• Về mặt Kỹ thuật: Thiết kế và triển khai thành công kiến trúc 3 lớp phân tách (Decoupled Architecture), hệ thống RESTful API chuẩn hóa trên Laravel 11, xác thực an toàn bằng Sanctum Token, và cơ sở dữ liệu MySQL 8.0 được tối ưu hóa.", bullet=True)
    p("• Về mặt Nội dung Game: Tích hợp thành công kho game phong phú (hơn 26 tựa game) gồm game nhúng HTML5 hợp pháp và tựa game do chính tác giả phát triển bằng Unity Engine xuất bản WebGL.", bullet=True)
    p("• Về mặt Quản trị: Xây dựng trang Admin CMS độc lập, trực quan, hỗ trợ quản lý toàn diện các thực thể Game, Danh mục, Người dùng và Kiểm duyệt bình luận.", bullet=True)

    h2("2. Hạn chế của đề tài")
    p("Mặc dù đạt được những kết quả tích cực, đề tài vẫn tồn tại một số hạn chế nhất định do giới hạn về mặt thời gian:")
    p("• Chưa tích hợp hệ thống thanh toán trực tuyến thực tế (Payment Gateway) để nâng cấp tài khoản VIP.", bullet=True)
    p("• Chưa phát triển chế độ chơi game nhiều người (Multiplayer) qua giao thức Real-time WebSockets.", bullet=True)

    h2("3. Hướng phát triển trong tương lai")
    p("Trong thời gian tới, hệ thống có thể được tiếp tục nâng cấp và phát triển theo các hướng sau:")
    p("1. Xây dựng ứng dụng di động Native (React Native / Flutter) dùng chung bộ Backend RESTful API hiện tại.")
    p("2. Phát triển tính năng Multiplayer thời gian thực (Real-time Leaderboard, Đấu trường 1v1) sử dụng Laravel Reverb hoặc Socket.io.")
    p("3. Tích hợp thuật toán trí tuệ nhân tạo (AI Recommendation) gợi ý tựa game phù hợp dựa trên thói quen và lịch sử chơi của thành viên.")

    # ---------------------------------------------------------------------------
    # TÀI LIỆU THAM KHẢO
    # ---------------------------------------------------------------------------
    h1("TÀI LIỆU THAM KHẢO")

    refs = [
        "1. Taylor Otwell (2024), Laravel Documentation (Version 11.x), https://laravel.com/docs/11.x",
        "2. Meta Open Source (2024), React Framework Documentation, https://react.dev",
        "3. Unity Technologies (2023), Unity User Manual 2022.3 LTS - WebGL Publishing, https://docs.unity3d.com/Manual/webgl.html",
        "4. Robin Wieruch (2022), The Road to React: Your journey to master plain Web Development with React, Leanpub.",
        "5. Matt Stauffer (2019), Laravel: Up & Running: A Hands-On Guide to Programming Laravel (2nd Edition), O'Reilly Media.",
        "6. World Wide Web Consortium (W3C) (2023), HTML5 & WebGL 2.0 Specification Standard, https://www.w3.org/TR/html52/",
        "7. GameDistribution Developer Program (2024), HTML5 SDK Integration Guide, https://gamedistribution.com/sdk",
        "8. Martin Fowler (2002), Patterns of Enterprise Application Architecture, Addison-Wesley Professional."
    ]

    for ref in refs:
        p(ref, space_after=6)

    # Save Document
    output_filename = "Bao_Cao_Do_An_webgame.docx"
    try:
        doc.save(output_filename)
        print(f"Document successfully created and saved as '{output_filename}'! Size: {os.path.getsize(output_filename)} bytes")
    except PermissionError:
        alt_filename = "Bao_Cao_Do_An_webgame_New.docx"
        doc.save(alt_filename)
        print(f"Permission denied on '{output_filename}' (file is likely open in MS Word). Saved as '{alt_filename}' instead! Size: {os.path.getsize(alt_filename)} bytes")

if __name__ == "__main__":
    build_document()
